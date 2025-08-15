import streamlit as st
from docx import Document
from docx.shared import Inches
import os
import io
import subprocess
import tempfile

def docx_to_pdf_libreoffice(docx_bytes, docx_filename):
    """Konversi DOCX ke PDF menggunakan LibreOffice"""
    try:
        # Buat direktori temp
        temp_dir = tempfile.mkdtemp()
        
        # Simpan DOCX ke temp file
        docx_path = os.path.join(temp_dir, docx_filename)
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        
        # Konversi ke PDF
        cmd = [
            "libreoffice", 
            "--headless", 
            "--convert-to", 
            "pdf", 
            "--outdir", 
            temp_dir, 
            docx_path
        ]
        
        subprocess.run(cmd, check=True)
        
        # Baca hasil PDF
        pdf_path = os.path.join(temp_dir, docx_filename.replace(".docx", ".pdf"))
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()
        
        # Hapus file temporary
        os.remove(docx_path)
        os.remove(pdf_path)
        os.rmdir(temp_dir)
        
        return pdf_bytes
    
    except Exception as e:
        st.error(f"Gagal mengkonversi ke PDF: {str(e)}")
        return None

def generate_cv(data, photo_path):
    doc = Document()

    # Header dengan foto
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    cell1 = row.cells[0]
    cell2 = row.cells[1]

    if photo_path and os.path.exists(photo_path):
        try:
            cell1.paragraphs[0].add_run().add_picture(photo_path, width=Inches(1.5))
        except Exception as e:
            st.warning(f"Tidak dapat menambahkan foto: {str(e)}")

    cell2.add_paragraph(data['name']).bold = True
    cell2.add_paragraph(f"{data['phone']} | {data['email']}")
    cell2.add_paragraph(f"{data['address']}")

    doc.add_paragraph()

    # Profil
    doc.add_heading("Profil Profesional", level=1)
    doc.add_paragraph(data['summary'])

    # Pengalaman Kerja
    doc.add_heading("Pengalaman Kerja", level=1)
    for exp in data['experience']:
        doc.add_heading(f"{exp['job_title']} – {exp['company']}", level=2)
        doc.add_paragraph(f"{exp['location']} | {exp['start_date']} – {exp['end_date']}")
        for task in exp['tasks']:
            doc.add_paragraph(f"• {task}", style='List Bullet')

    # Pendidikan
    doc.add_heading("Pendidikan", level=1)
    for edu in data['education']:
        doc.add_paragraph(f"{edu['degree']} – {edu['school']} ({edu['year']}) | IPK: {edu['gpa']}")

    # Pengalaman Organisasi
    doc.add_heading("Pengalaman Organisasi", level=1)
    for org in data['organizations']:
        doc.add_heading(f"{org['role']} – {org['name']}", level=2)
        doc.add_paragraph(f"{org['location']} | {org['start_date']} – {org['end_date']}")
        for desc in org['description']:
            doc.add_paragraph(f"• {desc}", style='List Bullet')

    # Keahlian
    doc.add_heading("Kemampuan", level=1)
    doc.add_paragraph(f"Hard Skills: {', '.join(data['hard_skills'])}")
    doc.add_paragraph(f"Soft Skills: {', '.join(data['soft_skills'])}")

    # Pencapaian
    doc.add_heading("Pencapaian Lainnya", level=1)
    for ach in data['achievements']:
        doc.add_paragraph(f"• {ach}", style='List Bullet')

    # Simpan ke buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def main():
    st.set_page_config(page_title="CV Builder", layout="wide")
    st.title("📄 CV Builder")

    # Informasi Pribadi
    st.header("📋 Informasi Pribadi")
    col1, col2 = st.columns(2)

    with col1:
        name = st.text_input("Nama Lengkap", "")
        phone = st.text_input("Nomor Telepon", "")

    with col2:
        email = st.text_input("Email", "")
        address = st.text_area("Alamat", "")

    # Upload Foto
    photo_file = st.file_uploader("📷 Upload Foto (Opsional)", type=["jpg", "jpeg", "png"])
    photo_path = None
    if photo_file:
        photo_path = f"temp_photo_{photo_file.name}"
        with open(photo_path, "wb") as f:
            f.write(photo_file.getbuffer())

    # Profil
    st.header("👤 Profil Profesional")
    summary = st.text_area("Ringkasan Profil", "")

    # Pengalaman Kerja
    st.header("💼 Pengalaman Kerja")
    experience = []
    exp_count = st.number_input("Jumlah pengalaman kerja", min_value=0, max_value=10, value=0)

    for i in range(exp_count):
        with st.expander(f"Pengalaman #{i+1}"):
            col1, col2 = st.columns(2)
            with col1:
                company = st.text_input(f"Perusahaan", key=f"company_{i}")
                job_title = st.text_input(f"Jabatan", key=f"job_title_{i}")
                location_exp = st.text_input(f"Lokasi", key=f"location_exp_{i}")
            with col2:
                start_date = st.text_input(f"Tanggal Mulai", key=f"start_date_{i}")
                end_date = st.text_input(f"Tanggal Selesai", key=f"end_date_{i}")
            
            tasks = st.text_area(f"Tugas/Pencapaian (pisahkan dengan koma)", key=f"tasks_{i}").split(",")
            experience.append({
                "company": company,
                "job_title": job_title,
                "location": location_exp,
                "start_date": start_date,
                "end_date": end_date,
                "tasks": [t.strip() for t in tasks if t.strip()]
            })

    # Pendidikan
    st.header("🎓 Pendidikan")
    education = []
    edu_count = st.number_input("Jumlah pendidikan", min_value=0, max_value=5, value=0)

    for i in range(edu_count):
        with st.expander(f"Pendidikan #{i+1}"):
            col1, col2 = st.columns(2)
            with col1:
                school = st.text_input(f"Sekolah/Kampus", key=f"school_{i}")
                degree = st.text_input(f"Gelar/Jurusan", key=f"degree_{i}")
            with col2:
                year = st.text_input(f"Tahun Lulus", key=f"year_{i}")
                gpa = st.text_input(f"IPK", key=f"gpa_{i}")
            
            education.append({"school": school, "degree": degree, "year": year, "gpa": gpa})

    # Organisasi
    st.header("🏢 Pengalaman Organisasi")
    organizations = []
    org_count = st.number_input("Jumlah organisasi", min_value=0, max_value=10, value=0)

    for i in range(org_count):
        with st.expander(f"Organisasi #{i+1}"):
            col1, col2 = st.columns(2)
            with col1:
                name_org = st.text_input(f"Nama Organisasi", key=f"name_org_{i}")
                role_org = st.text_input(f"Peran", key=f"role_org_{i}")
                location_org = st.text_input(f"Lokasi", key=f"location_org_{i}")
            with col2:
                start_org = st.text_input(f"Tanggal Mulai", key=f"start_org_{i}")
                end_org = st.text_input(f"Tanggal Selesai", key=f"end_org_{i}")
            
            desc_org = st.text_area(f"Deskripsi (pisahkan dengan koma)", key=f"desc_org_{i}").split(",")
            organizations.append({
                "name": name_org,
                "role": role_org,
                "location": location_org,
                "start_date": start_org,
                "end_date": end_org,
                "description": [d.strip() for d in desc_org if d.strip()]
            })

    # Keahlian
    st.header("🛠️ Kemampuan")
    col1, col2 = st.columns(2)

    with col1:
        hard_skills = st.text_area("Hard Skills (pisahkan dengan koma)", "").split(",")

    with col2:
        soft_skills = st.text_area("Soft Skills (pisahkan dengan koma)", "").split(",")

    # Pencapaian
    st.header("🏆 Pencapaian Lainnya")
    achievements = st.text_area("Pencapaian (pisahkan dengan koma)", "").split(",")

    st.markdown("---")
    if st.button("🚀 Generate CV", type="primary"):
        if not name or not phone or not email:
            st.error("Mohon isi Nama, Nomor Telepon, dan Email")
            return

        data = {
            "name": name,
            "phone": phone,
            "email": email,
            "address": address,
            "summary": summary,
            "experience": experience,
            "education": education,
            "organizations": organizations,
            "hard_skills": [x.strip() for x in hard_skills if x.strip()],
            "soft_skills": [x.strip() for x in soft_skills if x.strip()],
            "achievements": [x.strip() for x in achievements if x.strip()]
        }

        with st.spinner("Membuat CV..."):
            # Generate DOCX
            docx_bytes = generate_cv(data, photo_path)
            
            # Buat nama file
            safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            docx_filename = f"CV_{safe_name.replace(' ', '_')}.docx"
            pdf_filename = docx_filename.replace(".docx", ".pdf")
            
            # Konversi ke PDF menggunakan LibreOffice
            pdf_bytes = docx_to_pdf_libreoffice(docx_bytes, docx_filename)
            
            if pdf_bytes:
                st.success("CV berhasil dibuat!")
                
                # Tombol download PDF
                st.download_button(
                    label="📥 Download PDF",
                    data=pdf_bytes,
                    file_name=pdf_filename,
                    mime="application/pdf"
                )
            
            # Tombol download DOCX (siapapun)
            st.download_button(
                label="📥 Download DOCX",
                data=docx_bytes,
                file_name=docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # Clean up
    if photo_path and os.path.exists(photo_path):
        try:
            os.remove(photo_path)
        except:
            pass

if __name__ == "__main__":
    main()
